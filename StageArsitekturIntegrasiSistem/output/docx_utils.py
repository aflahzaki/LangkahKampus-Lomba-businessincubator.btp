"""
Shared utility module for generating TOGAF Enterprise Architecture DOCX documents.
Provides consistent formatting, cover pages, TOC, references, and helper functions.
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# Constants
FONT_NAME = "Times New Roman"
FONT_SIZE_MAIN = Pt(12)
FONT_SIZE_TITLE = Pt(14)
FONT_SIZE_SUBTITLE = Pt(12)
FONT_SIZE_FOOTNOTE = Pt(10)
LINE_SPACING = 1.5
MARGIN = Cm(2.54)  # 1 inch

LOGO_PATH = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
    "Logo LangkahKampus.png"
)

TEAM_MEMBERS = [
    ("Aflah Rafilah Zaki", "1301223259"),
    ("Azka Fathir Syarif", "1301220297"),
    ("Daffa Rizky Herdiawan", "1301223086"),
    ("Muhammad Arifin Ilham", "1301223300"),
]

DOSEN = "Dr. Farisya Setiadi ST., MTI"
ORGANIZATION = "LangkahKampus"
UNIVERSITY = "Telkom University"
COURSE = "Arsitektur Integrasi Sistem"


def create_document():
    """Create a new document with proper formatting settings."""
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = MARGIN
        section.bottom_margin = MARGIN
        section.left_margin = MARGIN
        section.right_margin = MARGIN

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE_MAIN
    style.paragraph_format.line_spacing = LINE_SPACING

    # Set East Asian font
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}" w:eastAsia="{FONT_NAME}" w:cs="{FONT_NAME}"/>')
        rpr.append(rFonts)
    else:
        rFonts.set(qn('w:ascii'), FONT_NAME)
        rFonts.set(qn('w:hAnsi'), FONT_NAME)
        rFonts.set(qn('w:eastAsia'), FONT_NAME)
        rFonts.set(qn('w:cs'), FONT_NAME)

    # Configure heading styles
    for i in range(1, 4):
        heading_style_name = f'Heading {i}'
        if heading_style_name in doc.styles:
            h_style = doc.styles[heading_style_name]
            h_style.font.name = FONT_NAME
            h_style.font.color.rgb = RGBColor(0, 0, 0)
            if i == 1:
                h_style.font.size = FONT_SIZE_TITLE
                h_style.font.bold = True
            elif i == 2:
                h_style.font.size = FONT_SIZE_SUBTITLE
                h_style.font.bold = True
            else:
                h_style.font.size = FONT_SIZE_MAIN
                h_style.font.bold = True
            h_style.paragraph_format.line_spacing = LINE_SPACING

            # Set heading font family
            h_rpr = h_style.element.get_or_add_rPr()
            h_rFonts = h_rpr.find(qn('w:rFonts'))
            if h_rFonts is None:
                h_rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}" w:eastAsia="{FONT_NAME}" w:cs="{FONT_NAME}"/>')
                h_rpr.append(h_rFonts)
            else:
                h_rFonts.set(qn('w:ascii'), FONT_NAME)
                h_rFonts.set(qn('w:hAnsi'), FONT_NAME)
                h_rFonts.set(qn('w:eastAsia'), FONT_NAME)
                h_rFonts.set(qn('w:cs'), FONT_NAME)

    return doc


def add_cover_page(doc, title, stage_name):
    """Add a cover page with logo, title, stage name, team, and organization."""
    # Add some spacing at the top
    for _ in range(2):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0

    # Add logo
    if os.path.exists(LOGO_PATH):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(LOGO_PATH, width=Cm(5))
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("[Logo LangkahKampus]")
        run.font.size = FONT_SIZE_MAIN

    # Spacing
    doc.add_paragraph()

    # Course name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TUGAS BESAR")
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_TITLE
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(COURSE.upper())
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_TITLE
    run.font.bold = True

    # Spacing
    doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title.upper())
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_TITLE
    run.font.bold = True

    # Stage name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(stage_name)
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_SUBTITLE
    run.font.bold = True

    # Spacing
    doc.add_paragraph()

    # Organization
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Studi Kasus: {ORGANIZATION}")
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_MAIN

    # Spacing
    doc.add_paragraph()

    # Team members
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Disusun Oleh:")
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_MAIN
    run.font.bold = True

    doc.add_paragraph()

    for name, nim in TEAM_MEMBERS:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{name} ({nim})")
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE_MAIN

    # Spacing
    doc.add_paragraph()

    # Dosen
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Dosen Pengampu: {DOSEN}")
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_MAIN

    # Spacing
    doc.add_paragraph()

    # University and year
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(UNIVERSITY.upper())
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_MAIN
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("2024")
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_MAIN
    run.font.bold = True

    # Page break after cover
    doc.add_page_break()


def add_table_of_contents(doc, sections):
    """
    Add a Daftar Isi (Table of Contents) page.
    sections: list of tuples (section_title, subsections_list)
    subsections_list is a list of subsection titles.
    """
    add_heading(doc, "DAFTAR ISI", level=1, centered=True)
    doc.add_paragraph()

    counter = 1
    for section_title, subsections in sections:
        # Main section entry
        p = doc.add_paragraph()
        run = p.add_run(f"{counter}. {section_title}")
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE_MAIN
        run.font.bold = True
        p.paragraph_format.line_spacing = LINE_SPACING

        # Subsections
        if subsections:
            sub_counter = 1
            for sub in subsections:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                run = p.add_run(f"{counter}.{sub_counter} {sub}")
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE_MAIN
                p.paragraph_format.line_spacing = LINE_SPACING
                sub_counter += 1

        counter += 1

    # Page break after TOC
    doc.add_page_break()


def add_heading(doc, text, level=1, centered=False):
    """Add a heading with proper formatting."""
    heading = doc.add_heading(text, level=level)
    if centered:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Ensure font settings
    for run in heading.runs:
        run.font.name = FONT_NAME
        run.font.color.rgb = RGBColor(0, 0, 0)
        if level == 1:
            run.font.size = FONT_SIZE_TITLE
            run.font.bold = True
        elif level == 2:
            run.font.size = FONT_SIZE_SUBTITLE
            run.font.bold = True
        else:
            run.font.size = FONT_SIZE_MAIN
            run.font.bold = True

        # Set font family on the run
        rpr = run._element.get_or_add_rPr()
        rFonts = rpr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}" w:eastAsia="{FONT_NAME}" w:cs="{FONT_NAME}"/>')
            rpr.append(rFonts)
        else:
            rFonts.set(qn('w:ascii'), FONT_NAME)
            rFonts.set(qn('w:hAnsi'), FONT_NAME)
            rFonts.set(qn('w:eastAsia'), FONT_NAME)
            rFonts.set(qn('w:cs'), FONT_NAME)

    return heading


def add_paragraph(doc, text, bold=False, italic=False, indent=None):
    """Add a paragraph with proper formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_MAIN
    run.font.bold = bold
    run.font.italic = italic
    p.paragraph_format.line_spacing = LINE_SPACING
    if indent:
        p.paragraph_format.left_indent = Cm(indent)

    # Justify text
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_bullet_list(doc, items, indent=1.27):
    """Add bullet point list items."""
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(indent)
        p.paragraph_format.first_line_indent = Cm(-0.63)
        run = p.add_run(f"\u2022  {item}")
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE_MAIN
        p.paragraph_format.line_spacing = LINE_SPACING


def add_numbered_list(doc, items, start=1, indent=1.27):
    """Add numbered list items."""
    for i, item in enumerate(items, start=start):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(indent)
        p.paragraph_format.first_line_indent = Cm(-0.63)
        run = p.add_run(f"{i}.  {item}")
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE_MAIN
        p.paragraph_format.line_spacing = LINE_SPACING


def add_table(doc, headers, rows, title=None, table_number=None):
    """Add a formatted table with optional title/number."""
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label = f"Tabel {table_number}. " if table_number else ""
        run = p.add_run(f"{label}{title}")
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE_MAIN
        run.font.bold = True
        p.paragraph_format.line_spacing = LINE_SPACING

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE_MAIN
        run.font.bold = True
        # Shade header
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)

    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = ""
            p = row_cells[i].paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.name = FONT_NAME
            run.font.size = FONT_SIZE_MAIN

    doc.add_paragraph()  # spacing after table
    return table


def add_figure_caption(doc, caption, figure_number):
    """Add a figure caption below an image."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Gambar {figure_number}. {caption}")
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_MAIN
    run.font.italic = True
    p.paragraph_format.line_spacing = LINE_SPACING
    return p


def add_references(doc, references):
    """
    Add Daftar Pustaka (References) page in APA format.
    references: list of reference strings in APA format.
    """
    doc.add_page_break()
    add_heading(doc, "DAFTAR PUSTAKA", level=1, centered=True)
    doc.add_paragraph()

    for ref in references:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)
        run = p.add_run(ref)
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE_MAIN
        p.paragraph_format.line_spacing = LINE_SPACING


def save_document(doc, filename):
    """Save the document to the output directory."""
    output_dir = os.path.dirname(os.path.abspath(__file__))
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    print(f"Document saved: {filepath}")
    return filepath
